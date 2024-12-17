import docx


def delete_first_page(doc_path):
    """
    Delete content from the first page of a Word document, stopping once likely to be past the first page.

    Args:
        doc_path (str): Path to the document.
    """
    doc = docx.Document(f'{doc_path}')
    elements_removed = 0  # Count of elements removed
    max_removal_limit = 1  # A safety net to avoid deleting the whole document

    # Remove paragraphs until content on the first page is cleared
    while len(doc.paragraphs) > 0 and elements_removed < max_removal_limit:
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

        elements_removed += 1

        # After removing some elements, check the new first paragraph
        if elements_removed >= 5 and len(doc.paragraphs) > 0:
            text_after_removal = doc.paragraphs[0].text.strip()
            if text_after_removal:
                break  # Likely the start of content from the second page

    # Optionally, remove any tables on the first page
    while len(doc.tables) > 0 and elements_removed < max_removal_limit:
        doc.tables[0]._element.getparent().remove(doc.tables[0]._element)
        elements_removed += 1

        if elements_removed >= 5 and len(doc.paragraphs) > 0:
            break  # Stop after clearing what seems to be the first page's content

    # Save the modified document
    doc.save(f'{doc_path}')

