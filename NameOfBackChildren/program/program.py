import shutil
import sys
import os


# Add the parent directory of 'program' to sys.path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))


from docx_copy_noBreak.docx_copy_noBreak import copy_layout_noBreak
import docx
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import mysql.connector as c
from docx.shared import Pt
import tkinter as tk
from tkinter import ttk
from tkinter import messagebox
import threading
from ttkthemes import ThemedStyle
from PIL import Image, ImageTk

# ========================================================================================================================================================================
mydb = c.connect(host="localhost", user="root", password="root", database="test")
mycursor = mydb.cursor()


def ensure_index(table, index_name, columns):
    check_index_query = f"""
    SELECT COUNT(*)
    FROM information_schema.STATISTICS
    WHERE table_schema = '{mydb.database}'
    AND table_name = '{table}'
    AND index_name = '{index_name}';
    """
    mycursor.execute(check_index_query)
    index_exists = mycursor.fetchone()[0] > 0

    if index_exists:
        print(f"Index '{index_name}' already exists on table '{table}'.")

    else:
        print(f"Index '{index_name}' does not exist on table '{table}', creating it now.")

        create_index_query = f"CREATE INDEX {index_name} ON {table}({columns});"
        try:
            mycursor.execute(create_index_query)
            mydb.commit()
            print(f"Index '{index_name}' created successfully.")

        except c.errors.ProgrammingError as e:
            print(f"Failed to create index '{index_name}' on table '{table}'. Error: {e}")
        except Exception as e:
            print(f"An unexpected error occurred while creating index '{index_name}' on table '{table}'. Error: {e}")

# Index definitions
indexes = [
    {"table": "student_marks_grade", "index_name": "idx_student_marks_grade", "columns": "student_scheme_subject_id, student_id, session_id, subject_id"},
    {"table": "student_scheme_subject", "index_name": "idx_student_scheme_subject", "columns": "id, student_id, subject_id"},
    {"table": "student_scheme_subject", "index_name": "idx_student_scheme_subject", "columns": "id, subject_id"},
    {"table": "subject", "index_name": "idx_subject", "columns": "id"},
    {"table": "class_registration", "index_name": "idx_class_registration", "columns": "student_id, scheme_id, session_id"},
    {"table": "scheme", "index_name": "idx_scheme", "columns": "id"},
    {"table": "student_marks_grade", "index_name": "idx_student_marks_grade", "columns": "student_scheme_subject_id, session_id"},
    {"table": "subject", "index_name": "idx_subject_id", "columns": "id"},
    {"table": "student", "index_name": "idx_student_id", "columns": "id"},
    {"table": "scheme", "index_name": "idx_scheme_term_id", "columns": "term_id, id"},

    # New indexes from previous query
    {"table": "student_scheme_subject", "index_name": "idx_student_scheme_subject_student_subject", "columns": "id, student_id, subject_id"},
    {"table": "student_marks_grade", "index_name": "idx_student_marks_grade_scheme_subject_session", "columns": "student_scheme_subject_id, session_id"},
    {"table": "subject", "index_name": "idx_subject_id", "columns": "id"},
    {"table": "student", "index_name": "idx_student_id", "columns": "id"},
    {"table": "class_registration", "index_name": "idx_class_registration_student_scheme_session", "columns": "student_id, scheme_id, session_id"},
    {"table": "scheme", "index_name": "idx_scheme_id_term_session", "columns": "id, term_id, session_id"},

    # Newly added indexes
    {"table": "grade_card", "index_name": "idx_grade_card_student_session_term", "columns": "student_id, session_id, term_id"},
    {"table": "student", "index_name": "idx_student_id_year_term", "columns": "id, year_term_code"}
]

# Ensure indexes are present
for index in indexes:
    ensure_index(index['table'], index['index_name'], index['columns'])

running = False


def on_submit():
    # Get the input from the entry widget

    if entry.get() == "":
        messagebox.showerror("ERROR!!", 'Enter a code!!')
    else:
        global running
        text_area.delete("1.0", tk.END)
        user_input = entry.get()
        session_input = year_code_entry.get()
        if session_input == "":
            session_input = 22
        print("the session code is :" ,session_input)
        print(user_input)
        running = True
        threading.Thread(target=report_generation, args=(user_input,session_input), daemon=True).start()


# Define the path to the original .docx file
original_file = 'head.docx'


def remove_empty_list(lst):
    return [sublist for sublist in lst if sublist]


def calculate_CGPA(lst):
    try:
        CGPA = 0
        for i in lst:
            CGPA = CGPA + i
        return CGPA/len(lst)
    except ZeroDivisionError as e:
        return


# query_grade_result = ("""
# SELECT theory_credit, practical_credit, theory_grade_id, practical_grade_id,session_id, subject_id, subject.name
# FROM test.student_marks_grade
# JOIN (SELECT * FROM test.student_scheme_subject) AS studentSubject
# JOIN test.subject
# WHERE student_marks_grade.student_scheme_subject_id = studentSubject.id AND
# student_id=42 AND
# subject.id=studentSubject.subject_id AND
# session_id > 7 AND
# subject_id IN (59, 60, 61, 62 ,63 ,64 ,65) AND
# studentSubject.student_id IN (SELECT id FROM test.student
# WHERE id IN (SELECT student_id FROM test.class_registration
# WHERE scheme_id IN (SELECT id FROM test.scheme
# WHERE term_id=90) AND session_id < 16 AND session_id > 7))
# ORDER BY student_id ASC, session_id ASC, subject_id ASC;
# """)
# mycursor.execute(query_grade_result)
# grade_result = mycursor.fetchall()
# mydb.commit()

# query_session_id = ("SELECT distinct session_id "
#                     "FROM test.student_marks_grade "
#                     "JOIN (SELECT * FROM test.student_scheme_subject) AS studentSubject "
#                     "JOIN test.subject "
#                     "WHERE student_marks_grade.student_scheme_subject_id = studentSubject.id AND "
#                     "student_id=42 AND "
#                     "subject.id=studentSubject.subject_id AND "
#                     f"session_id > {user_input[0]-8} AND "
#                     "studentSubject.student_id IN (SELECT id FROM test.student "
#                     "WHERE id IN (SELECT student_id FROM test.class_registration "
#                     "WHERE scheme_id IN (SELECT id FROM test.scheme "
#                     f"WHERE term_id=90) AND session_id < {user_input[0]+1} AND session_id > {user_input[0]-8})) "
#                     "ORDER BY session_id ASC;")
# mycursor.execute(query_session_id)
# session_id = mycursor.fetchall()
# mydb.commit()


def make_specific_paragraphs_bold(doc):
    for para in doc.paragraphs:
        if "Enrol." in para.text:
            for run in para.runs:
                run.bold = True


def process_document(input_path):
    # Load the document
    doc = Document(f"{input_path}")

    # Make specific paragraphs bold
    make_specific_paragraphs_bold(doc)

    # Save the document with the changes
    doc.save(f"{input_path}")


def make_specific_paragraphs_bold2(doc):
    for para in doc.paragraphs:
        if "RollNo" in para.text:
            for run in para.runs:
                run.bold = True


def process_document2(input_path):
    # Load the document
    doc = Document(f"{input_path}")

    # Make specific paragraphs bold
    make_specific_paragraphs_bold2(doc)

    # Save the document with the changes
    doc.save(f"{input_path}")


def set_line_spacing(file, spacing):
    """
    Set the line spacing for all paragraphs in a document, including those inside tables.

    Args:
        file (str): Path to the .docx file.
        spacing (float): Line spacing to be set. (e.g., 1.0, 1.5, 2.0)
    """
    doc = docx.Document(f'{file}')

    # Adjust line spacing for paragraphs not in tables
    for paragraph in doc.paragraphs:
        para_format = paragraph.paragraph_format
        para_format.line_spacing = Pt(spacing * 12)  # Convert spacing to points

    # Adjust line spacing for paragraphs inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    para_format = paragraph.paragraph_format
                    para_format.line_spacing = Pt(spacing * 12)  # Convert spacing to points

    output_path = f'{file}'

    if os.path.exists(output_path):
        os.remove(output_path)  # Remove the final file if it exists

    doc.save(output_path)  # Save the document to the correct file path


def on_item_selected(event):
    # Get the selected item
    selected_item = batch_details.selection()
    if selected_item:
        # Get the values of the selected item
        item_values = batch_details.item(selected_item[0], 'values')
        # Populate the entry boxes with the selected item values
        entry.delete(0, tk.END)
        entry.insert(0, item_values[0])
def on_item_selected2(event):
    # Get the selected item
    selected_item2 = batch_details2.selection()
    if selected_item2:
        # Get the values of the selected item
        item_values2 = batch_details2.item(selected_item2[0], 'values')
        # Populate the entry boxes with the selected item values
        year_code_entry.delete(0, tk.END)
        year_code_entry.insert(0, item_values2[1])


def insert_table_after_paragraph(paragraph, data, cell_widths=None):
    # Create a new table element
    table = OxmlElement('w:tbl')

    # Create table properties
    tbl_props = OxmlElement('w:tblPr')
    table.append(tbl_props)

    # Create table grid
    tbl_grid = OxmlElement('w:tblGrid')
    for _ in range(len(data[0])):
        tbl_grid.append(OxmlElement('w:gridCol'))
    table.append(tbl_grid)

    # Add rows and cells to the table
    for row_index, row_data in enumerate(data):
        row = OxmlElement('w:tr')
        for col_index, cell_data in enumerate(row_data):
            cell = OxmlElement('w:tc')

            # Set cell width if specified
            if cell_widths and col_index < len(cell_widths):
                cell_width = cell_widths[col_index]
                cell_props = OxmlElement('w:tcPr')
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(cell_width))  # Set the width (in twips)
                tcW.set(qn('w:type'), 'dxa')  # Width type 'dxa' for twips
                cell_props.append(tcW)
                cell.append(cell_props)

            cell_text = OxmlElement('w:p')
            run = OxmlElement('w:r')
            text = OxmlElement('w:t')
            text.text = str(cell_data)
            run.append(text)
            cell_text.append(run)
            cell.append(cell_text)
            row.append(cell)

        table.append(row)

    # Insert the table directly after the given paragraph
    parent = paragraph._element.getparent()
    parent.insert(parent.index(paragraph._element) + 1, table)


def replace_text_in_docx(file, old_text, new_text, data):
    # Load the existing document
    doc = Document(f'{file}')

    for para in doc.paragraphs:
        if old_text in para.text:
            para.text = para.text.replace(old_text, new_text)

    # Handle the special placeholder for the table
    paragraphs_to_remove = []
    for para in doc.paragraphs:
        if '[Subject Detail]' in para.text:
            # Replace placeholder with a marker
            para.text = para.text.replace('[Subject Detail]', '[TABLE_PLACEHOLDER]')

            # Insert the table right after the paragraph with the placeholder
            new_data = data
            # Specify cell widths (in twips, 1 inch = 1440 twips)
            cell_widths = [1700, 4000, 1500, 1500]

            insert_table_after_paragraph(para, new_data, cell_widths)

            # Remove the placeholder marker text
            para.text = para.text.replace('[TABLE_PLACEHOLDER]', '')

            # Optionally add text from the next paragraph if needed
            paragraphs_to_remove.append(para)  # To remove any unwanted extra paragraphs

    # Remove unwanted paragraphs if needed
    for para in paragraphs_to_remove:
        if para.text == '':
            para._element.getparent().remove(para._element)

    # Save the modified document
    new_file = f"temp_{old_text.replace(' ', '_')}.docx"
    doc.save(f'{new_file}')
    if file != original_file:
        os.remove(f'{file}')
    return f'{new_file}'


# Initialize the dictionary
session_to_year = {}

# Mapping session values to their corresponding years
mycursor.execute("SELECT DISTINCT year FROM scheme;")
years = mycursor.fetchall()

for year in years:
    year_value = year[0]
    # Calculate the mapped value based on the pattern
    if year_value >= 2015:
        mapped_value = 15 + 2 * (year_value - 2015)
        session_to_year[mapped_value] = str(year_value)

print("session_to_year : ", session_to_year)

# Fetch session names and IDs
mycursor.execute("SELECT id, name FROM test.session;")
session_names = mycursor.fetchall()

# Create a dictionary with name as the key and id as the value
session_dict = {name: id for id, name in session_names}

data_old = []

# Print the dictionary
print(session_dict) 

data_old = []


def stop_task():
    global running
    running = False
    submit_button['state'] = 'normal'


def report_generation(BatchCode, session_input):
    global running
    while running:
        if not running:
            break
        user_input = [BatchCode]

        # Fetch all the name and group data
        query_student_grp_details = (f"""
                SELECT student_id, student_branch_details.name , grp1, grp2, grp3, grp4, grp5, grp6, grp7, grp8, enroll_no, app_fname, app_lname, father_name, mother_name, branch_mode, roll_no, roll_no_2, roll_no_3, roll_no_4, roll_no_5, roll_no_6, roll_no_7, roll_no_8
                FROM test.stgrp 
                JOIN (SELECT * FROM test.student
                WHERE id IN (SELECT distinct student_id 
                FROM (
                    SELECT * 
                    FROM test.grade_card 
                    WHERE student_id IN (
                        SELECT id 
                        FROM test.student 
                        WHERE year_term_code = "{str(user_input[0])}"
                        AND id in (SELECT distinct student_id 
                                FROM test.class_registration 
                                WHERE session_id =  {session_input}
                                )))
                                AS term_student 
                WHERE student_id IN (
                    SELECT student_id 
                    FROM test.grade_card 
                    WHERE student_id IN (
                        SELECT id 
						FROM test.student 
						WHERE year_term_code = "{str(user_input[0])}"
						AND id in (SELECT distinct student_id 
								   FROM test.class_registration 
								   WHERE session_id =  {session_input}
						)
                )))) AS session_4th_yr 
                JOIN (SELECT code, name FROM test.branch) AS student_branch_details
                WHERE stgrp.student_id = session_4th_yr.id AND
                code = session_4th_yr.branch_code
                ORDER BY student_id ASC;
                """)

        mycursor.execute(query_student_grp_details)
        student_grp_details = mycursor.fetchall()

        # Fetch all unique subjects for all students once
        unique_subjects_query = (f"""
SELECT 
    us.student_id, 
    us.subject_id, 
    us.session_id
FROM 
    (
        SELECT 
            studentSubject.subject_id,
            student_marks_grade.session_id,
            studentSubject.student_id
        FROM 
            test.student_marks_grade 
        JOIN 
            test.student_scheme_subject AS studentSubject 
            ON student_marks_grade.student_scheme_subject_id = studentSubject.id 
        JOIN 
            test.subject 
            ON subject.id = studentSubject.subject_id
        WHERE
            studentSubject.student_id IN (
                SELECT distinct student_id 
                FROM (
                    SELECT * 
                    FROM test.grade_card 
                    WHERE student_id IN (
                        SELECT id 
                        FROM test.student 
                        WHERE year_term_code = "{str(user_input[0])}"
                        AND id in (SELECT distinct student_id 
                                FROM test.class_registration 
                                WHERE session_id = {session_input}
)))
                                AS term_student 
                WHERE student_id IN (
                    SELECT student_id 
                    FROM test.grade_card 
                    WHERE student_id IN (
                        SELECT id 
						FROM test.student 
						WHERE year_term_code = "{str(user_input[0])}"
						AND id in (SELECT distinct student_id 
								   FROM test.class_registration 
								   WHERE session_id = {session_input}

						)
                ))
            )
    ) us
JOIN 
    (
        SELECT student_id, subject_id, MIN(session_id) AS min_session
        FROM 
        (
            SELECT 
                studentSubject.subject_id,
                student_marks_grade.session_id,
                studentSubject.student_id
            FROM 
                test.student_marks_grade 
            JOIN 
                test.student_scheme_subject AS studentSubject 
                ON student_marks_grade.student_scheme_subject_id = studentSubject.id 
            JOIN 
                test.subject 
                ON subject.id = studentSubject.subject_id
            WHERE
                studentSubject.student_id IN (
                    SELECT distinct student_id 
                FROM (
                    SELECT * 
                    FROM test.grade_card 
                    WHERE student_id IN (
                        SELECT id 
                        FROM test.student 
                        WHERE year_term_code = "{str(user_input[0])}"
                        AND id in (SELECT distinct student_id 
                                FROM test.class_registration 
                                WHERE session_id = {session_input}
)))
                                AS term_student 
                WHERE student_id IN (
                    SELECT student_id 
                    FROM test.grade_card 
                    WHERE student_id IN (
                        SELECT id 
						FROM test.student 
						WHERE year_term_code = "{str(user_input[0])}"
						AND id in (SELECT distinct student_id 
								   FROM test.class_registration 
								   WHERE session_id = {session_input}

						)
                ))
                )
        ) AS unique_subjects
        GROUP BY student_id, subject_id
    ) min_subjects 
ON us.student_id = min_subjects.student_id 
AND us.subject_id = min_subjects.subject_id 
AND us.session_id = min_subjects.min_session
ORDER BY 
    us.student_id ASC,
    us.session_id ASC, 
    us.subject_id ASC;
                """)
        mycursor.execute(unique_subjects_query)
        unique_subjects = mycursor.fetchall()

        submit_button['state'] = 'disabled'

        query_session_mnth = f"SELECT name FROM test.session WHERE id = {session_input};"
        mycursor.execute(query_session_mnth)
        session_mnth = mycursor.fetchall()
        mydb.commit()


        if unique_subjects == []:
            root.after(0, text_area.insert("end", "No students\nCOMPLETED!!"))
            root.update_idletasks()
            text_area.see(tk.END)
            progress['value'] = 100
            running = False
            break

        query_student_id = (f"""SELECT id, enroll_no FROM test.student
                WHERE id IN (SELECT distinct student_id 
                FROM (
                    SELECT * 
                    FROM test.grade_card 
                    WHERE student_id IN (
                        SELECT id 
                        FROM test.student 
                        WHERE year_term_code = "{str(user_input[0])}"
                        AND id in (SELECT distinct student_id 
                                FROM test.class_registration 
                                WHERE session_id = {session_input}
)))
                                AS term_student 
                WHERE student_id IN (
                    SELECT student_id 
                    FROM test.grade_card 
                    WHERE student_id IN (
                        SELECT id 
						FROM test.student 
						WHERE year_term_code = "{str(user_input[0])}"
						AND id in (SELECT distinct student_id 
								   FROM test.class_registration 
								   WHERE session_id = {session_input}

						)
                )))
                ORDER BY id ASC;""")
        mycursor.execute(query_student_id)
        student_ids = mycursor.fetchall()
        mydb.commit()

        delete_Variable = 1
        # Loop through the user input to set the year dynamically
        for session_value in [session_input]:
            if not running:
                break
            print(session_value)
            year_value2 = first_dict.get(str(user_input[0]), 'DefaultYear')  # Get the corresponding year or a default value
            year_value = session_to_year.get(int(session_input), 'DefaultYear')  # Get the corresponding year or a default value

            # Dictionary of old and new text pairs
            # the value will be changed as requested

            print(f"BackStudent{year_value2}Batch.docx")

            # Define the path to the original .docx file
            original_file = 'head.docx'
            # Rename the final document if needed
            final_file = f'BackStudent{year_value2}Batch.docx'
            # Use the copy_layout function
            newDoc = docx.Document()
            newDoc.save(f'BackStudent{year_value2}Batch.docx')
            newdoc = f'BackStudent{year_value2}Batch.docx'

            current_file = newdoc
            copy_layout_noBreak(current_file, "resultSem.docx")

            subjects_by_student = {}
            for student_id, subject_id, session_id in unique_subjects:
                if student_id not in subjects_by_student:
                    subjects_by_student[student_id] = {}
                if session_id not in subjects_by_student[student_id]:
                    subjects_by_student[student_id][session_id] = []
                subjects_by_student[student_id][session_id].append(subject_id)

            # Access and print each semester of every student individually
            studentNo = 1
            for student_id, sessions in subjects_by_student.items():
                if not running:
                    break
                copy_layout_noBreak(current_file, original_file)

                replacements = {}
                student_detail = []
                for i in student_grp_details:
                    if student_id == i[0]:
                        replacements = {
                            '[En_Roll]': str(i[10]),
                            '[Name]': i[11] + " " + i[12],
                            '[Branch]': str(i[1]),
                            '[last session]': str(year_value2),
                            '[year]': year_value,
                            '[FName]': i[13],
                            '[MName]': i[14],
                            '[Regular]': i[15]
                        }
                        # print(i)
                        root.after(0, text_area.insert("end", f"{i}\n"))
                        root.update_idletasks()
                        text_area.see(tk.END)
                        for j in i:
                            student_detail.append('-' if j == '' else j)

                # Perform replacements
                for old_text, new_text in replacements.items():
                    current_file = replace_text_in_docx(current_file, old_text, new_text, '')

                process_document(current_file)

                # print(f"Student ID: {student_id}")
                studentProgressLabel.config(text=f"Student: {studentNo}/{len(student_ids)}")
                studentRollNo.config(text=f"Student En. Roll No.: {student_ids[studentNo-1][1]}")
                percentage = (studentNo/len(student_ids))*100
                progress['value'] = percentage
                root.update_idletasks()
                root.after(0, text_area.insert("end", f"Student ID: {student_id}\n"))
                root.update_idletasks()
                text_area.see(tk.END)

                if studentNo == len(subjects_by_student.items()):
                    root.update_idletasks()
                    root.after(0, text_area.insert("end", "COMPLETED!"))
                    root.update_idletasks()
                    running = False
                    break

                studentNo += 1

            set_line_spacing(current_file, .80)

            if os.path.exists(f'{final_file}'):
                os.remove(f'{final_file}')  # Remove the final file if it exists
            os.rename(f'{current_file}', f'{final_file}')
            shutil.move(final_file, f'../../Reports/BackStudents/{final_file}')

            progress['value'] = 100
            root.update_idletasks()
            studentProgressLabel.config(text="Student: --")
            studentRollNo.config(text="Student Roll No.: --")

            submit_button['state'] = 'normal'


# Set up the main application window
root = tk.Tk()
root.title("4th Year Report")
root.config(bg="#D3D3D3")
root.attributes('-fullscreen', True)
root.tk.call("tk", "scaling", 1.5)

frame = tk.Frame(root)
frame.place(x=0, y=100, height=400, width=500)
frame.config(bg="#D3D3D3")

Symbol_Frame = tk.Frame(root, relief='groove', borderwidth=3)
Symbol_Frame.place(x=1245, y=20, height=40, width=78)

MinimizeLogo = Image.open("./all_.jpg")
MinimizeLogo = MinimizeLogo.resize((30, 30))
MinimizeLogo = ImageTk.PhotoImage(MinimizeLogo)

CloseLogo = Image.open("./close_.jpg")
CloseLogo = CloseLogo.resize((30, 30))
CloseLogo = ImageTk.PhotoImage(CloseLogo)

Button_MinimizeLogo = tk.Button(Symbol_Frame, command=root.iconify, cursor="hand2", image=MinimizeLogo)
Button_MinimizeLogo.pack(side="left")

Button_CloseLogo = tk.Button(Symbol_Frame, command=root.destroy, cursor="hand2", image=CloseLogo)
Button_CloseLogo.pack(side="left")

label_head = ttk.Label(root, text="BACK STUDENT REPORT GENERATION")
label_head.pack(anchor='n', pady=20)

label_head.config(font=("Times New Roman", 20))

batchFrames = tk.Frame(root, bg="red")
batchFrames.place(x=460, y=100, height=150, width=300)

batchFrames2 = tk.Frame(root, bg="red")
batchFrames2.place(x=760, y=100, height=150, width=300)

batch_details = ttk.Treeview(batchFrames, columns=("Year", "Session Code"), show='headings')

vsb = tk.Scrollbar(batchFrames, orient="vertical", command=batch_details.yview)
batch_details.config(yscrollcommand=vsb.set)

# Define the column headings
batch_details.heading("Year", text="Year")
batch_details.heading("Session Code", text="Session Code")

batch_details.column("Year", width=50)
batch_details.column("Session Code", width=50)

# Execute the first query and fetch results
mycursor.execute("SELECT DISTINCT year FROM scheme;")
years = mycursor.fetchall()

# Execute the second query and fetch results
mycursor.execute("SELECT distinct year_term_code FROM test.student;")
year_term = mycursor.fetchall()

# Create a dictionary from the first and second queries
first_dict = {term[0]: year[0] for year, term in zip(years, year_term)}

print("first_dict :" ,first_dict)

# Create a new dictionary where the values from the first dictionary are used as keys
second_dict = {key: value for key, value in first_dict.items()}

# Create the new dictionary
new_dict = {}

# Iterate over the second dictionary to build the new dictionary
for key, year in second_dict.items():
    # Find the corresponding value from the first dictionary where the year is the value
    for mapped_value, mapped_year in session_to_year.items():
        if int(mapped_year) == year:  # Match the year as an integer
            new_dict[key] = mapped_value  # Set the key from second_dict and the value from first_dict

# Print the new dictionary
print("new_dict : ", new_dict)

for key, value in new_dict.items():
    batch_details.insert("", "end", values=(key, value))

batch_details.pack(side=tk.LEFT, fill='both', expand=True)
vsb.pack(side=tk.RIGHT, fill=tk.Y)

batch_details.bind("<<TreeviewSelect>>", on_item_selected)

batch_details2 = ttk.Treeview(batchFrames2, columns=("Session", "Session Code"), show='headings')

vsb = tk.Scrollbar(batchFrames2, orient="vertical", command=batch_details.yview)
batch_details.config(yscrollcommand=vsb.set)

# Define the column headings
batch_details2.heading("Session", text="Session")
batch_details2.heading("Session Code", text="Session Code")

batch_details2.column("Session", width=50)
batch_details2.column("Session Code", width=50)

mycursor.execute("SELECT DISTINCT year FROM scheme;")
years = mycursor.fetchall()

# Insert values from the dictionary into the Treeview
for name, id in session_dict.items():
    batch_details2.insert("", "end", values=(name, id))

batch_details2.pack(side=tk.LEFT, fill='both', expand=True)
vsb.pack(side=tk.RIGHT, fill=tk.Y)

batch_details2.bind("<<TreeviewSelect>>", on_item_selected2)

# Create a style object
style = ttk.Style()

themeStyle = ThemedStyle(root)
# Define custom styles for the Treeview
themeStyle.theme_use("winnative")  # Use the "clam" theme for better styling options

# Style the Treeview background, foreground, and font
style.configure("Treeview",
                background="#D3D3D3",
                foreground="black",
                rowheight=25,
                fieldbackground="#D3D3D3",
                font=("Fixedsys", 12))

# Style the Treeview heading
style.configure("Treeview.Heading",
                background="#4CAF50",
                foreground="white",
                font=("Fixedsys", 12, "bold"))

# Style the selected item
style.map("Treeview",
          background=[('selected', '#347083')],
          foreground=[('selected', 'white')])

frame_detail = tk.Frame(root)
frame_detail.place(x=1080, y=100, height=100, width=400)
frame_detail.config(bg="#D3D3D3")

label_name = ttk.Label(frame_detail, text="Created By: Chinmay Gupta\n\t    DE22103\n\t    BE 3rd yr Computer Engg\n\t    CSA")
label_name.pack(anchor='w')

label_name.config(font=("Fixedsys", 11))

# Create a label for the entry widget
label = tk.Label(frame, text="Year Term Code:", font=("Fixedsys", 6))
label.pack(padx=100, pady=5, anchor='w')

# Create an entry widget that only accepts integer input
entry = ttk.Entry(frame)
entry.pack(padx=100, pady=5, anchor='w')

label2 = tk.Label(frame, text="Session Code:", font=("Fixedsys", 6))
label2.pack(padx=100, pady=5, anchor='w')

year_code_entry = ttk.Entry(frame)
year_code_entry.pack(padx=100, pady=5, anchor='w')

# Create a submit button
submit_button = tk.Button(frame, text="Submit", command=on_submit, relief="ridge", font=("Fixedsys", 3), bg="#4CAF50", fg="white")
submit_button.place(x=270, y=101)

# Create a button to stop the long-running task
stop_button = tk.Button(frame, text="Stop Task", command=stop_task, relief="ridge", font=("Fixedsys", 3), bg="darkred", fg="white")
stop_button.place(x=350, y=101)

frame_textArea = tk.Frame(root, bg="red")
frame_textArea.place(x=100, y=300, height=450, width=1300)

scrollbar = tk.Scrollbar(frame_textArea)
scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

studentProgressLabel = tk.Label(root, text="Student: --", font=("Fixedsys", 3))
studentProgressLabel.place(x=100, y=270)

studentRollNo = tk.Label(root, text="Student En. Roll No.: --", font=("Fixedsys", 3))
studentRollNo.place(x=300, y=270)

progressBarStyle = ttk.Style()
progressBarStyle.configure("green.Horizontal.TProgressbar",
                troughcolor='#696969',  # Color of the bar's background
                background='green',       # Color of the bar's progress
                thickness=20)

progress = ttk.Progressbar(root, orient='horizontal', length=1285, mode='determinate', style="green.Horizontal.TProgressbar")
progress.place(x=100, y=765)

text_area = tk.Text(frame_textArea)
text_area.pack(fill='both', expand=True)

scrollbar.config(command=text_area.yview)

# Start the main event loop
root.mainloop()

mycursor.close()
mydb.commit()
