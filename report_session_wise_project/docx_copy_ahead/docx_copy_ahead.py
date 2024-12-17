import docx


def copy_layout_ahead(existing_doc_path, document_copy_path):
    """
    Copy the layout from the existing document and add it ahead of the original document.

    Args:
        document_copy_path (str): Path to the copying document.
        existing_doc_path (str): Path to the existing document.
    """
    existing_doc = docx.Document(f'{existing_doc_path}')
    document_copy = docx.Document(f'{document_copy_path}')

    # Create a new document object, but don't create a new file
    new_doc = existing_doc

    # Add a page break before copying the content
    new_doc.add_page_break()

    # Iterate over the paragraphs and tables in the original order
    for paragraph in document_copy.paragraphs:
        new_paragraph = new_doc.add_paragraph()
        new_paragraph.text = paragraph.text
        new_paragraph.paragraph_format.alignment = paragraph.paragraph_format.alignment
        new_paragraph.paragraph_format.space_before = paragraph.paragraph_format.space_before
        new_paragraph.paragraph_format.space_after = paragraph.paragraph_format.space_after

    for table in document_copy.tables:
        new_table = new_doc.add_table(rows=table.rows, cols=table.columns)
        new_table.style = table.style
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                new_table.cell(row_idx, col_idx).text = cell.text

    # Save the modified document
    new_doc.save(f'{existing_doc_path}')
