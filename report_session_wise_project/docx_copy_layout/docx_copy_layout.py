import docx
import os


def copy_layout(original_file, dest_file):
    # Open the existing document
    existing_doc = docx.Document(f'{original_file}')

    # Create a new document
    new_doc = docx.Document()

    # Copy the layout from the existing document to the new document
    for paragraph in existing_doc.paragraphs:
        new_paragraph = new_doc.add_paragraph()
        new_paragraph.text = paragraph.text
        new_paragraph.paragraph_format.alignment = paragraph.paragraph_format.alignment
        new_paragraph.paragraph_format.space_before = paragraph.paragraph_format.space_before
        new_paragraph.paragraph_format.space_after = paragraph.paragraph_format.space_after

    for table in existing_doc.tables:
        new_table = new_doc.add_table(rows=table.rows, cols=table.columns)
        new_table.style = table.style
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                new_table.cell(row_idx, col_idx).text = cell.text

    if os.path.exists(f'{dest_file}'):
        os.remove(f'{dest_file}')  # Remove the final file if it exists
    # Save the new document
    new_doc.save(f'{dest_file}')
